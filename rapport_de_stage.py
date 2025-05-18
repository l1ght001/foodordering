from docx import Document
from docx.shared import Pt

# Create a new Document
doc = Document()

# Title
doc.add_heading('Internship Report: foodislice', 0)

def add_section(title, content):
    doc.add_heading(title, level=1)
    for paragraph in content:
        doc.add_paragraph(paragraph)

# Sections content
data = [
    ("Introduction", [
        "During my internship, I had the opportunity to work on a web-based food ordering and management platform called foodislice. The project aimed to provide restaurants and food businesses with a modern, user-friendly admin panel to manage their menu, orders, and customer interactions efficiently."
    ]),
    ("Company/Project Presentation", [
        "foodislice is a digital solution designed for restaurants and food vendors to streamline their operations. The platform offers an intuitive admin dashboard, menu management, order tracking, and payment configuration, all accessible through a web interface. The project was developed as part of a team effort to address the growing need for digital transformation in the food service industry."
    ]),
    ("Project Objectives", [
        "- Develop a responsive and modern admin panel for food business management.",
        "- Enable easy menu and category management.",
        "- Provide real-time order tracking and status updates.",
        "- Allow configuration of payment and delivery settings.",
        "- Ensure a secure authentication system for administrators."
    ]),
    ("Technical Stack", [
        "- Frontend: React.js, Material-UI, Framer Motion",
        "- Backend: Node.js, Express.js",
        "- Database: PostgreSQL (via Prisma ORM)",
        "- Other Tools: Axios, TailwindCSS, JWT for authentication"
    ]),
    ("Features Developed", [
        "- Dashboard: Overview of orders, earnings, customers, and products.",
        "- Menu Management: Add, edit, and remove menu items and categories.",
        "- Order Management: Track and update order statuses (pending, completed, rejected).",
        "- Payment Settings: Configure currency, service fee, and delivery fee.",
        "- Settings: Customize display options and manage admin credentials.",
        "- Export/Invoice: Export orders as CSV and generate invoices in PDF format."
    ]),
    ("My Contributions", [
        "- Implemented the menu management interface with category toggles and item CRUD operations.",
        "- Developed the dashboard with real-time statistics and recent orders table.",
        "- Integrated authentication and security settings for admin users.",
        "- Enhanced the UI/UX using Material-UI and custom themes.",
        "- Added features for exporting data and generating invoices."
    ]),
    ("Challenges and Solutions", [
        "- Challenge: Handling real-time updates for orders and menu changes.\n  Solution: Used React state management and localStorage for persistence.",
        "- Challenge: Ensuring secure authentication.\n  Solution: Implemented JWT-based authentication and secure password management.",
        "- Challenge: Creating a responsive and intuitive UI.\n  Solution: Leveraged Material-UI components and custom themes for consistency."
    ]),
    ("Skills Acquired", [
        "- Advanced React.js development and state management.",
        "- UI/UX design with Material-UI and TailwindCSS.",
        "- Backend API integration and authentication.",
        "- Database management with Prisma and PostgreSQL.",
        "- Team collaboration and agile development practices."
    ]),
    ("Conclusion", [
        "This internship allowed me to apply and expand my technical skills in a real-world project. I gained valuable experience in full-stack development, problem-solving, and working within a team. The foodislice project is a testament to the impact of digital solutions in the food industry, and I am proud to have contributed to its development."
    ]),
]

for title, content in data:
    add_section(title, content)

# Save the document
output_path = 'rapport_de_stage_foodislice.docx'
doc.save(output_path)
print(f"Rapport de stage generated: {output_path}") 